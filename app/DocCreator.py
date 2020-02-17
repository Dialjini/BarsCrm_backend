from app import app, models, db
from flask import send_from_directory
from datetime import datetime, timedelta
from docx import Document, styles
from docx.shared import Pt
import os
import json
from num2t4ru import num2text


def replace_request(words, replacements, doc):
    for j in range(0, len(words)):
        if not replacements[j]:
            replacements[j] = ' '
        for i in doc.tables[0]._cells:
            if words[j] in i.text:
                i.text = i.text.replace(str(words[j]), str(replacements[j]))
    return doc


def replace_doc(words, replacements, doc):
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"
    for j in range(0, len(words)):
        if not replacements[j]:
            replacements[j] = ' '

        for i in doc.paragraphs:
            if words[j] in i.text:
                i.style = style
                i.text = i.text.replace(str(words[j]), str(replacements[j]))
        try:
            if words[j] in doc.tables[0].cell(0, 1).text:
                doc.tables[0].cell(0, 1).text = doc.tables[0].cell(0, 1).text.replace(str(words[j]), str(replacements[j]))
        except Exception:
            continue
    return doc


def Generate_Transit(dir_u, date, delivery, adress, adress2, type):
    document = models.Document()
    delivery = models.Delivery.query.filter_by(id=delivery).first()
    carrier = models.Carrier.query.filter_by(id=delivery.Carrier_id).first()
    document.MonthNum = models.getMonthNum()
    document.Date = str(datetime.now().month) + '/' + str(datetime.now().year)
    document.Creation_date = str(datetime.now().day) + '.' + str(datetime.now().month) + '.' + str(datetime.now().year)
    transit_type = adress.split(' ')
    print('ok')
    Sum = 0
    for i in json.loads(delivery.Amounts):
        Sum += Sum + float(str(i['sum']).replace(' ', ''))
    if int(Sum) == Sum:
        Sum = int(Sum)
    else:
        Sum = round(Sum, 2)


    print(Sum)
    print('| sum')
    document.Client_name = 'transit'

    Items = models.Item.query.all()
    item_info = {'mass': 0}
    for i in Items:
        if str(i.Item_id) in json.loads(delivery.Item_ids):
            item_info['mass'] = str(float(item_info['mass']) + float(i.Weight))
            item_info['packing'] = i.Packing

    db.session.add(document)
    db.session.commit()
    document.Path = '{}.docx'.format('TransitN' + str(document.id))
    db.session.commit()

    if transit_type[0] == 'OOO':
        doc = Document(os.path.dirname(__file__) + '/files/Zayavka_OOO.docx')
    else:
        doc = Document(os.path.dirname(__file__) + '/files/Zayavka_IP.docx')

    if not delivery.Load_type:
        delivery.Load_type = ' '

    if transit_type[1] == 'OOO':
        doc = replace_doc(doc=doc, words=['document.Date', 'date.d', 'date.m', 'date.y',
                                    'document.Client_contact_name', 'document.Client_name'],
                    replacements=[document.Date, date.d, date.m, date.y, document.Client_contact_name, 'ООО "Барс"'])
        doc = replace_request(words=['{{Имя_клиента}}', '{{ИНН}}', '{{КПП}}', '{{Юр_адрес}}',
                                     '{{МаркаАвто}}', '{{Контакт водителя}}', '{{ПаспортныеДанныеВодителя}}',
                                     '{{Адресс_погрузки}}', '{{Контактное_лицо}}', '{{Наименование_грузополучателя}}',
                                     '{{Адрес_разгрузки}}', '{{Дата_и_время_разгрузки}}', '{{Контактное лицо на выгрузке}}',
                                     '{{Вес_груза}}', '{{Вид_упаковки}}', '{{Способ погрузки}}', '{{Погрузка}}',
                                     '{{Разгрузка}}', '{{Сумма}}', '{{Сумма_словами}}'],
                              replacements=[carrier.Name, carrier.UHH, carrier.Bik,
                                            carrier.Address, delivery.Auto,
                                            delivery.Contact_Name, delivery.Passport_data,
                                            delivery.Contact_Number, delivery.Contact_End, 'ООО "Барс"', adress2,
                                            delivery.End_date, delivery.Contact_Start.replace('-s!s-', ','), str(item_info['mass']),
                                            item_info['packing'], delivery.Load_type, delivery.Date, delivery.End_date,
                                            str(Sum), num2text(Sum)], doc=doc)
    else:
        doc = replace_doc(doc=doc, words=['document.Date', 'date.d', 'date.m', 'date.y',
                                    'document.Client_contact_name', 'document.Client_name'],
                    replacements=[document.Date, date.d, date.m, date.y, document.Client_contact_name, 'ООО "Барс"'])
        doc = replace_request(words=['{{Имя_клиента}}', '{{ИНН}}', '{{КПП}}', '{{Юр_адрес}}',
                                     '{{МаркаАвто}}', '{{Контакт водителя}}', '{{ПаспортныеДанныеВодителя}}',
                                     '{{Адресс_погрузки}}', '{{Контактное_лицо}}', '{{Наименование_грузополучателя}}',
                                     '{{Адрес_разгрузки}}', '{{Дата_и_время_разгрузки}}', '{{Контактное лицо на выгрузке}}',
                                     '{{Вес_груза}}', '{{Вид_упаковки}}', '{{Способ погрузки}}', '{{Погрузка}}',
                                     '{{Разгрузка}}', '{{Сумма}}', '{{Сумма_словами}}'],
                              replacements=[carrier.Name, carrier.UHH, carrier.Bik,
                                            carrier.Address, delivery.Auto,
                                            delivery.Contact_Name, delivery.Passport_data,
                                            delivery.Contact_Number, delivery.Contact_End, 'ИП Балкина Ирина Николаевна', adress2,
                                            delivery.End_date, delivery.Contact_Start.replace('-s!s-', ','), str(item_info['mass']),
                                            item_info['packing'], delivery.Load_type, delivery.Date, delivery.End_date,
                                            str(Sum), num2text(Sum)], doc=doc)

    doc.save(dir_u + '/{}.docx'.format('TransitN' + str(document.id)))
    return send_from_directory(directory=os.path.abspath(os.path.dirname(__file__) + '/upload'),
                                   filename=document.Path)

def Generate_DogovorNaDostavkuOOO(dir_u, owner, date):
    document = models.Document()
    document.MonthNum = models.getMonthNum()
    document.Date = str(datetime.now().month) + '/' + str(datetime.now().year)
    document.Creation_date = str(datetime.now().day) + '.' + str(datetime.now().month) + '.' + str(datetime.now().year)

    document.Prefix = 'ООО'
    document.UHH = owner.UHH
    document.Owner_type = owner.__tablename__
    if owner.Bik:
        document.Bik = owner.Bik
    else:
        document.Bik = ''
    document.KPP = owner.kpp
    if owner.rc:
        document.rc = owner.rc
    else:
        document.rc = ''
    if owner.kc:
        document.kc = owner.kc
    else:
        document.kc = ''

    document.Client_contact_name = owner.Director
    document.Owner_id = owner.id
    try:
        document.Client_prefix_address = owner.Adress
    except Exception:
        document.Client_prefix_address = owner.Address
    document.Client_name = owner.Name
    try:
        document.Client_mail_address = owner.Adress
    except Exception:
        document.Client_mail_address = owner.Address

    db.session.add(document)
    db.session.commit()
    document.Path = '{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id))
    db.session.commit()
    rate = datetime.now() + timedelta(days=7)
    doc = Document(os.path.dirname(__file__) + '/files/Dogovor_mezhdu_OOO_Bars_i_perevozchikom.docx')
    doc = replace_doc(doc=doc, words=['document.MonthNum', 'document.Date', 'date.d', 'date.m', 'date.y',
                                      'document.Client_contact_name', 'rate.day', 'rate.month', 'rate.year',
                                      'document.Client_name', 'document.Client_prefix_address',
                                      'document.Client_mail_address', 'document.UHH', 'document.KPP', 'document.rc',
                                      'document.kc', 'document.Bik'],
                      replacements=[document.MonthNum, document.Date, date.d, date.m, date.y,
                                    document.Client_contact_name,
                                    rate.day, rate.month, rate.year, document.Client_name,
                                    document.Client_prefix_address, document.Client_mail_address, document.UHH,
                                    document.KPP,
                                    document.rc, document.kc, document.Bik])
    doc.save(dir_u + '/{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id)))
    return 'OK'



def Generate_DogovorNaDostavkuIP(dir_u, owner, date):
    document = models.Document()
    document.MonthNum = models.getMonthNum()
    document.Date = str(datetime.now().month) + '/' + str(datetime.now().year)
    document.Creation_date = str(datetime.now().day) + '.' + str(datetime.now().month) + '.' + str(datetime.now().year)

    document.Prefix = 'ИП'
    document.UHH = owner.UHH
    document.Owner_type = owner.__tablename__
    if owner.Bik:
        document.Bik = owner.Bik
    else:
        document.Bik = ''
    document.KPP = owner.kpp
    if owner.rc:
        document.rc = owner.rc
    else:
        document.rc = ''
    if owner.kc:
        document.kc = owner.kc
    else:
        document.kc = ''

    document.Client_contact_name = owner.Director
    document.Owner_id = owner.id
    try:
        document.Client_prefix_address = owner.Adress
    except Exception:
        document.Client_prefix_address = owner.Address
    document.Client_name = owner.Name
    try:
        document.Client_mail_address = owner.Adress
    except Exception:
        document.Client_mail_address = owner.Address

    db.session.add(document)
    db.session.commit()
    document.Path = '{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id))
    db.session.commit()
    rate = datetime.now() + timedelta(days=7)
    doc = Document(os.path.dirname(__file__) + '/files/DOGOVOR_mezhdu_IP_i_perevozchikom.docx')
    doc = replace_doc(doc=doc, words=['document.MonthNum', 'document.Date', 'date.d', 'date.m', 'date.y',
                                      'document.Client_contact_name', 'rate.day', 'rate.month', 'rate.year',
                                      'document.Client_name', 'document.Client_prefix_address',
                                      'document.Client_mail_address', 'document.UHH', 'document.KPP', 'document.rc',
                                      'document.kc', 'document.Bik'],
                      replacements=[document.MonthNum, document.Date, date.d, date.m, date.y,
                                    document.Client_contact_name,
                                    rate.day, rate.month, rate.year, document.Client_name,
                                    document.Client_prefix_address, document.Client_mail_address, document.UHH,
                                    document.KPP,
                                    document.rc, document.kc, document.Bik])
    doc.save(dir_u + '/{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id)))
    return 'OK'


def Generate_Dogovor_na_tovari_ooo(dir_u, owner, date):
    document = models.Document()
    document.MonthNum = models.getMonthNum()
    document.Date = str(datetime.now().month) + '/' + str(datetime.now().year)
    document.Creation_date = str(datetime.now().day) + '.' + str(datetime.now().month) + '.' + str(datetime.now().year)

    document.Prefix = 'ООО'
    document.UHH = owner.UHH
    document.Owner_type = owner.__tablename__
    if owner.Bik:
        document.Bik = owner.Bik
    else:
        document.Bik = ''
    document.KPP = owner.kpp
    if owner.rc:
        document.rc = owner.rc
    else:
        document.rc = ''
    if owner.kc:
        document.kc = owner.kc
    else:
        document.kc = ''

    document.Client_contact_name = owner.Director
    document.Owner_id = owner.id
    try:
        document.Client_prefix_address = owner.Adress
    except Exception:
        document.Client_prefix_address = owner.Address
    document.Client_name = owner.Name
    try:
        document.Client_mail_address = owner.Adress
    except Exception:
        document.Client_mail_address = owner.Address

    db.session.add(document)
    db.session.commit()
    document.Path = '{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id))
    db.session.commit()

    rate = datetime.now() + timedelta(days=7)
    doc = Document(os.path.dirname(__file__) + '/files/Образец_Договора_на_товары.docx')
    doc = replace_doc(doc=doc, words=['document.MonthNum', 'document.Date', 'date.d', 'date.m', 'date.y',
                                'document.Client_contact_name', 'rate.day', 'rate.month', 'rate.year',
                                'document.Client_name', 'document.Client_prefix_address',
                                'document.Client_mail_address', 'document.UHH', 'document.KPP', 'document.rc',
                                'document.kc', 'document.Bik'],
                replacements=[document.MonthNum, document.Date, date.d, date.m, date.y, document.Client_contact_name,
                              rate.day, rate.month, rate.year, document.Client_name,
                              document.Client_prefix_address, document.Client_mail_address, document.UHH, document.KPP,
                              document.rc, document.kc, document.Bik])
    doc.save(dir_u + '/{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id)))
    return 'OK'


def Generate_Dogovor_na_tovari_ip(dir_u, owner, date):
    document = models.Document()
    document.MonthNum = models.getMonthNum()
    document.Date = str(datetime.now().month) + '/' + str(datetime.now().year)
    document.Creation_date = str(datetime.now().day) + '.' + str(datetime.now().month) + '.' + str(datetime.now().year)
    document.Prefix = 'ИП'

    document.Owner_type = owner.__tablename__
    document.UHH = owner.UHH
    if owner.Bik:
        document.Bik = owner.Bik
    else:
        document.Bik = ''
    document.KPP = owner.kpp
    if owner.rc:
        document.rc = owner.rc
    else:
        document.rc = ''
    if owner.kc:
        document.kc = owner.kc
    else:
        document.kc = ''

    document.Client_contact_name = owner.Director
    document.Owner_id = owner.id
    try:
        document.Client_prefix_address = owner.Adress
    except Exception:
        document.Client_prefix_address = owner.Address
    document.Client_name = owner.Name
    try:
        document.Client_mail_address = owner.Adress
    except Exception:
        document.Client_mail_address = owner.Address

    db.session.add(document)
    db.session.commit()
    document.Path = '{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id))
    db.session.commit()
    rate = datetime.now() + timedelta(days=7)
    doc = Document(os.path.dirname(__file__) + '/files/ДОГОВОР орион.docx')
    doc = replace_doc(doc=doc, words=['document.MonthNum', 'document.Date', 'date.d', 'date.m', 'date.y',
                                'document.Client_contact_name', 'rate.day', 'rate.month', 'rate.year',
                                'document.Client_name', 'document.Client_prefix_address',
                                'document.Client_mail_address', 'document.UHH', 'document.KPP', 'document.rc',
                                'document.kc', 'document.Bik'],
                replacements=[document.MonthNum, document.Date, date.d, date.m, date.y, document.Client_contact_name,
                              rate.day, rate.month, rate.year, document.Client_name,
                              document.Client_prefix_address, document.Client_mail_address, document.UHH, document.KPP,
                              document.rc, document.kc, document.Bik])
    doc.save(dir_u + '/{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id)))
    return 'OK'


def Generate_Zayavka_OOO(dir_u, owner, date, delivery):
    document = models.Document()
    delivery = models.Delivery.query.filter_by(id=delivery).first()
    Client = models.Client.query.filter_by(Name=delivery.Name).first()
    if not Client:
        Client = models.Provider.query.filter_by(Name=delivery.Name).first()
    if not Client:
        Client = models.Carrier.query.filter_by(Name=delivery.Name).first()

    document.MonthNum = models.getMonthNum()
    document.Date = str(datetime.now().month) + '/' + str(datetime.now().year)
    document.Creation_date = str(datetime.now().day) + '.' + str(datetime.now().month) + '.' + str(datetime.now().year)
    document.UHH = owner.UHH
    document.Owner_type = owner.__tablename__
    document.Prefix = 'ООО'
    Sum = 0
    item_info = {'mass': 0}
    for i in json.loads(delivery.Amounts):
        item_info['mass'] += float(str(i['volume']).replace(' ', ''))
        Sum += Sum + float(str(i['sum']).replace(' ', ''))
    if int(Sum) == Sum:
        Sum = int(Sum)
    else:
        Sum = round(Sum, 2)
    document.Bik = owner.Bik
    document.KPP = owner.kpp
    document.rc = owner.rc
    document.kc = owner.kc

    document.Client_contact_name = owner.Director
    document.Owner_id = owner.id
    try:
        document.Client_prefix_address = owner.Adress
    except Exception:
        document.Client_prefix_address = owner.Address
    document.Client_name = owner.Name
    try:
        document.Client_mail_address = owner.Adress
    except Exception:
        document.Client_mail_address = owner.Address
    Items = models.Item.query.all()
    item_info['packing'] = ''
    for i in Items:
        if str(i.Item_id) in json.loads(delivery.Item_ids):
            item_info['packing'] = i.Packing

    db.session.add(document)
    db.session.commit()
    document.Path = '{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id))
    db.session.commit()
    try:
        if Client.Fact_address:
            Adress = Client.Fact_address
        else:
            Adress = Client.Adress
    except Exception:
        if Client.__tablename__ == 'Carrier':
            Adress = Client.Address
        else:
            Adress = Client.Address

    doc = Document(os.path.dirname(__file__) + '/files/Zayavka_OOO.docx')
    doc = replace_doc(doc=doc, words=['document.Date', 'date.d', 'date.m', 'date.y',
                                'document.Client_contact_name', 'document.Client_name'],
                replacements=[document.Date, date.d, date.m, date.y, document.Client_contact_name, document.Client_name])
    doc = replace_request(words=['{{Имя_клиента}}', '{{ИНН}}', '{{КПП}}', '{{Юр_адрес}}',
                                 '{{МаркаАвто}}', '{{Контакт водителя}}', '{{ПаспортныеДанныеВодителя}}',
                                 '{{Адресс_погрузки}}', '{{Контактное_лицо}}', '{{Наименование_грузополучателя}}',
                                 '{{Адрес_разгрузки}}', '{{Дата_и_время_разгрузки}}', '{{Контактное лицо на выгрузке}}',
                                 '{{Вес_груза}}', '{{Вид_упаковки}}', '{{Способ погрузки}}', '{{Погрузка}}',
                                 '{{Разгрузка}}', '{{Сумма}}', '{{Сумма_словами}}'],
                          replacements=[document.Client_name, document.UHH, document.KPP,
                                        document.Client_prefix_address, delivery.Auto,
                                        delivery.Contact_Name + ', ' + delivery.Contact_Number, delivery.Passport_data,
                                        delivery.Stock.replace('-s!s-', ','), delivery.Contact_End, Client.Name, Adress.replace('-s!s-', ','),
                                        delivery.End_date, delivery.Contact_Start.replace('-s!s-', ','), str(item_info['mass']),
                                        item_info['packing'], delivery.Load_type, delivery.Date, delivery.End_date,
                                        str(Sum), num2text(Sum)], doc=doc)

    doc.save(dir_u + '/{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id)))
    return send_from_directory(directory=os.path.abspath(os.path.dirname(__file__) + '/upload'),
                                   filename=document.Path)



def Generate_Zayavka_IP(dir_u, owner, date, delivery):
    document = models.Document()
    delivery = models.Delivery.query.filter_by(id=delivery).first()
    Client = models.Client.query.filter_by(Name=delivery.Name).first()
    if not Client:
        Client = models.Provider.query.filter_by(Name=delivery.Name).first()
    if not Client:
        Client = models.Carrier.query.filter_by(Name=delivery.Name).first()
    document.MonthNum = models.getMonthNum()
    document.Date = str(datetime.now().month) + '/' + str(datetime.now().year)
    document.Creation_date = str(datetime.now().day) + '.' + str(datetime.now().month) + '.' + str(datetime.now().year)
    Sum = 0
    item_info = {'mass': 0}
    for i in json.loads(delivery.Amounts):
        item_info['mass'] += float(str(i['volume']).replace(' ', ''))
        Sum += Sum + float(str(i['sum']).replace(' ', ''))

    if int(Sum) == Sum:
        Sum = int(Sum)
    else:
        Sum = round(Sum, 2)
    document.UHH = owner.UHH
    document.Owner_type = owner.__tablename__
    document.Prefix = 'ИП'
    document.Bik = owner.Bik
    document.KPP = owner.kpp
    document.rc = owner.rc
    document.kc = owner.kc

    document.Client_contact_name = owner.Director
    document.Owner_id = owner.id
    try:
        document.Client_prefix_address = owner.Adress
    except Exception:
        document.Client_prefix_address = owner.Address
    document.Client_name = owner.Name
    try:
        document.Client_mail_address = owner.Adress
    except Exception:
        document.Client_mail_address = owner.Address

    Items = models.Item.query.all()
    item_info['packing'] = ''
    for i in Items:
        if str(i.Item_id) in json.loads(delivery.Item_ids):
            item_info['packing'] = i.Packing

    db.session.add(document)
    db.session.commit()
    document.Path = '{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id))
    db.session.commit()
    try:
        if Client.Fact_address:
            Adress = Client.Fact_address
        else:
            Adress = Client.Adress
    except Exception:
        if Client.__tablename__ == 'Carrier':
            Adress = Client.Address
        else:
            Adress = Client.Address

    doc = Document(os.path.dirname(__file__) + '/files/Zayavka_IP.docx')
    doc = replace_doc(doc=doc, words=['document.Date', 'date.d', 'date.m', 'date.y',
                                'document.Client_contact_name', 'document.Client_name'],
                replacements=[document.Date, date.d, date.m, date.y, document.Client_contact_name, document.Client_name])
    doc = replace_request(words=['{{Имя_клиента}}', '{{ИНН}}', '{{КПП}}', '{{Юр_адрес}}',
                                 '{{МаркаАвто}}', '{{Контакт водителя}}', '{{ПаспортныеДанныеВодителя}}',
                                 '{{Адресс_погрузки}}', '{{Контактное_лицо}}', '{{Наименование_грузополучателя}}',
                                 '{{Адрес_разгрузки}}', '{{Дата_и_время_разгрузки}}', '{{Контактное лицо на выгрузке}}',
                                 '{{Вес_груза}}', '{{Вид_упаковки}}', '{{Способ_погрузки}}', '{{Погрузка}}',
                                 '{{Разгрузка}}', '{{Сумма}}', '{{Сумма_словами}}'],
                          replacements=[document.Client_name, document.UHH, document.KPP,
                                        document.Client_prefix_address, delivery.Auto,
                                        delivery.Contact_Name + ', ' + delivery.Contact_Number, delivery.Passport_data,
                                        delivery.Stock.replace('-s!s-', ','), delivery.Contact_End, Client.Name, Adress.replace('-s!s-', ','),
                                        delivery.End_date, delivery.Contact_Start.replace('-s!s-', ','), str(item_info['mass']),
                                        item_info['packing'], delivery.Load_type, delivery.Date, delivery.End_date,
                                        str(Sum), num2text(Sum)], doc=doc)
    doc.save(dir_u + '/{}.docx'.format(owner.__tablename__ + str(owner.id) + 'N' + str(document.id)))
    return send_from_directory(directory=os.path.abspath(os.path.dirname(__file__) + '/upload'),
                                   filename=document.Path)
